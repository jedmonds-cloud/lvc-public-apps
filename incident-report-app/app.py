import streamlit as st
import requests
import io
import re
from datetime import datetime
from zoneinfo import ZoneInfo
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MELBOURNE_TZ = ZoneInfo("Australia/Melbourne")

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="LVC Incident Report Generator",
    page_icon="📋",
    layout="centered"
)

# ── LVC Brand colours ─────────────────────────────────────────────────────────
# Mid teal:    #3D9B9C  — headings, accents
# Brand teal:  #7ACCC8  — table headers
# Light grey:  #F2F2F2  — alternating rows

st.markdown("""
<style>
    .main { max-width: 820px; }
    .stApp { background-color: #f5f5f5; }

    /* Header banner */
    .report-header {
        background: #3D9B9C;
        color: white;
        padding: 22px 28px;
        border-radius: 8px;
        margin-bottom: 24px;
    }
    .report-header h1 { color: white; margin: 0; font-size: 22px; font-family: Calibri, sans-serif; }
    .report-header p  { color: #d4f0ef; margin: 5px 0 0 0; font-size: 14px; font-family: Calibri, sans-serif; }

    /* Status badges */
    .status-badge         { display: inline-block; padding: 4px 12px; border-radius: 4px; font-weight: bold; font-size: 13px; }
    .status-under-review  { background: #fff3cd; color: #856404; }
    .status-pending       { background: #d4f0ef; color: #1f6b6c; }
    .status-closed        { background: #d1e7dd; color: #0f5132; }
    .status-default       { background: #e2e3e5; color: #41464b; }

    /* Section headings */
    .section-title {
        color: #3D9B9C;
        font-weight: bold;
        font-size: 15px;
        font-family: Calibri, sans-serif;
        border-bottom: 2px solid #7ACCC8;
        padding-bottom: 5px;
        margin-top: 28px;
        margin-bottom: 12px;
    }

    /* Field rows */
    .field-row   { display: flex; padding: 8px 0; border-bottom: 1px solid #e9ecef; }
    .field-label { font-weight: 600; color: #3D9B9C; min-width: 210px; font-size: 14px; font-family: Calibri, sans-serif; }
    .field-value { color: #212529; font-size: 14px; flex: 1; font-family: Calibri, sans-serif; }

    /* Alternating field rows */
    .field-row:nth-child(even) { background-color: #F2F2F2; }

    /* Narrative boxes */
    .narrative-box {
        background: white;
        border: 1px solid #7ACCC8;
        border-left: 4px solid #3D9B9C;
        border-radius: 4px;
        padding: 14px 16px;
        margin: 8px 0 14px 0;
        font-size: 14px;
        line-height: 1.6;
        white-space: pre-wrap;
        font-family: Calibri, sans-serif;
    }
    .narrative-label {
        font-weight: 600;
        color: #3D9B9C;
        font-size: 13px;
        margin-bottom: 4px;
        font-family: Calibri, sans-serif;
        text-transform: uppercase;
        letter-spacing: 0.4px;
    }

    /* Warning box */
    .warning-box {
        background: #fff8e1;
        border: 1px solid #7ACCC8;
        border-radius: 6px;
        padding: 12px 16px;
        margin: 12px 0;
        font-size: 14px;
    }

    /* Button override */
    .stDownloadButton > button {
        background-color: #3D9B9C !important;
        color: white !important;
        border: none !important;
        font-family: Calibri, sans-serif !important;
    }
    .stDownloadButton > button:hover {
        background-color: #2e7b7c !important;
    }
    .stButton > button {
        background-color: #3D9B9C !important;
        color: white !important;
        border: none !important;
    }
</style>
""", unsafe_allow_html=True)

# ── Helpers ───────────────────────────────────────────────────────────────────
def get_api_key():
    try:
        return st.secrets["CLICKUP_API_KEY"]
    except:
        return None

def ts_to_date(ts_str):
    if not ts_str:
        return "—"
    try:
        ts = int(ts_str) / 1000
        dt = datetime.fromtimestamp(ts, tz=MELBOURNE_TZ)
        hour12 = int(dt.strftime("%I"))  # 1–12, no leading zero, portable
        ampm = dt.strftime("%p").lower()
        return f"{dt.day} {dt.strftime('%B %Y')}, {hour12}:{dt.strftime('%M')}{ampm}"
    except:
        return "—"

def ts_to_datetime(ts_str):
    if not ts_str:
        return None
    try:
        return datetime.fromtimestamp(int(ts_str) / 1000, tz=MELBOURNE_TZ)
    except:
        return None

def display_client(inc):
    client = (inc.get("client") or "").strip()
    if client in ("", "-", "—", "None"):
        return "Unnamed Client"
    return client

def build_filename(inc):
    client = (inc.get("client") or "").strip()
    task_id = (inc.get("id") or "").strip()
    date_short = (inc.get("report_date_short") or "").strip()

    if client in ("", "-", "—", "None"):
        name_part = f"Unnamed Client {task_id}".strip()
    else:
        name_part = client

    if date_short:
        base = f"Incident Report - {name_part} ({date_short})"
    else:
        base = f"Incident Report - {name_part}"

    # Replace characters Windows rejects with a space (keeps word boundaries)
    base = re.sub(r'[\\/:*?"<>|]', " ", base)
    # Collapse repeated spaces and underscores
    base = re.sub(r" {2,}", " ", base)
    base = re.sub(r"_{2,}", "_", base)
    base = base.strip()

    # Cap total length (including extension) at 150 characters
    ext = ".docx"
    if len(base) + len(ext) > 150:
        base = base[: 150 - len(ext)].rstrip()

    return base + ext

def get_field(fields, name):
    for f in fields:
        if f["name"] == name:
            val = f.get("value")
            if val is None:
                return None
            return str(val).strip() if str(val).strip() not in ["", "None"] else None
    return None

def get_dropdown_label(fields, name):
    for f in fields:
        if f["name"] == name:
            val = f.get("value")
            if val is None:
                return None
            try:
                idx = int(val)
                opts = f.get("type_config", {}).get("options", [])
                for o in opts:
                    if o.get("orderindex") == idx:
                        return o["name"]
            except:
                pass
    return None

def bool_display(fields, name):
    for f in fields:
        if f["name"] == name:
            val = f.get("value")
            if val in [True, "true", "True", 1, "1"]:
                return "Yes"
            elif val in [False, "false", "False", 0, "0"]:
                return "No"
    return "—"

def status_badge(status):
    s = status.lower()
    if "review" in s:
        cls = "status-under-review"
    elif "pending" in s:
        cls = "status-pending"
    elif "closed" in s or "complete" in s:
        cls = "status-closed"
    else:
        cls = "status-default"
    return f'<span class="status-badge {cls}">{status.upper()}</span>'

def fetch_incident(task_id, api_key):
    url = f"https://api.clickup.com/api/v2/task/{task_id.strip()}?custom_fields=true"
    headers = {"Authorization": api_key}
    r = requests.get(url, headers=headers, timeout=10)
    if r.status_code == 200:
        return r.json(), None
    elif r.status_code == 401:
        return None, "Invalid API key."
    elif r.status_code == 404:
        return None, f"Task ID '{task_id}' not found."
    else:
        return None, f"ClickUp API error {r.status_code}"

def parse_incident(data):
    cf = data.get("custom_fields", [])
    assignees = ", ".join([a["username"] for a in data.get("assignees", [])])
    lead_list = []
    for f in cf:
        if f["name"] == "Lead" and f.get("value"):
            lead_list = [u["username"] for u in f["value"]] if isinstance(f["value"], list) else []
    # Effective report date: Event Date & Time, falling back to date_created
    event_dt = ts_to_datetime(get_field(cf, "Event Date & Time")) or ts_to_datetime(data.get("date_created"))
    return {
        "id": data["id"],
        "status": data.get("status", {}).get("status", "Unknown").title(),
        "date_created": ts_to_date(data.get("date_created")),
        "event_datetime": ts_to_date(get_field(cf, "Event Date & Time")),
        "event_location": get_field(cf, "Event Location") or "—",
        "service": get_dropdown_label(cf, "Service") or "—",
        "client": get_field(cf, "Client Name") or "—",
        "staff": get_field(cf, "Staff Name") or "—",
        "other_people": get_field(cf, "Other People Involved") or "—",
        "assignees": assignees or "—",
        "lead": ", ".join(lead_list) if lead_list else "—",
        "review_due": ts_to_date(get_field(cf, "Date Review Due")),
        "incident_details": get_field(cf, "(1.) Incident Details") or "—",
        "impact": get_field(cf, "(2.) Impact on Client") or "—",
        "actions_taken": get_field(cf, "(3.) Immediate Actions Taken") or "—",
        "follow_up": get_field(cf, "(4.) Required Follow-Up Actions") or "—",
        "client_incident_type": get_dropdown_label(cf, "CLIENT INCIDENT TYPE") or "Not yet classified",
        "whs_incident_type": get_dropdown_label(cf, "WHS INCIDENT TYPE") or "Not applicable",
        "client_consulted": bool_display(cf, "Client Consulted"),
        "worker_consulted": bool_display(cf, "Worker Consulted"),
        "worker_injured": bool_display(cf, "Worker Injured"),
        "ndis_reportable": get_dropdown_label(cf, "NDIS Reportable") or "—",
        "worksafe_reportable": get_dropdown_label(cf, "Worksafe Reportable") or "—",
        "review_findings": get_field(cf, "Review Findings") or "—",
        "required_actions": get_field(cf, "Required Actions") or "—",
        "date_review_done": ts_to_date(get_field(cf, "Date LVC Review Done")),
        "date_ndis_notified": ts_to_date(get_field(cf, "Date NDIS Notified")),
        "ndis_ref": get_field(cf, "NDIS Reference No") or "—",
        "date_closed": ts_to_date(get_field(cf, "Date LVC Closed Incident")),
        "report_date_short": event_dt.strftime("%d%b%y") if event_dt else "",
        "report_date_long": f"{event_dt.day} {event_dt.strftime('%B %Y')}" if event_dt else "—",
    }

# ── Word doc generator ────────────────────────────────────────────────────────
# LVC colours: Mid teal #3D9B9C | Brand teal #7ACCC8 | Light grey #F2F2F2

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table_row(table, label, value, row_idx):
    row = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]
    label_cell.text = label
    value_cell.text = value or "—"
    # Label: teal text, light grey bg
    run = label_cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x3D, 0x9B, 0x9C)
    set_cell_bg(label_cell, "F2F2F2")
    # Value: alternating white / light grey
    vrun = value_cell.paragraphs[0].runs[0]
    vrun.font.size = Pt(10)
    vrun.font.name = "Calibri"
    if row_idx % 2 == 0:
        set_cell_bg(value_cell, "FFFFFF")
    else:
        set_cell_bg(value_cell, "F2F2F2")
    for cell in [label_cell, value_cell]:
        for par in cell.paragraphs:
            par.paragraph_format.space_before = Pt(3)
            par.paragraph_format.space_after = Pt(3)
            par.paragraph_format.left_indent = Cm(0.2)

def generate_word(inc):
    doc = Document()
    # Core document properties
    doc.core_properties.title = f"Incident Report - {display_client(inc)} ({inc['report_date_short']})"
    doc.core_properties.author = "La Vita Care"
    doc.core_properties.subject = inc["id"]
    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    run = title.add_run("INCIDENT REPORT")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x3D, 0x9B, 0x9C)
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.space_before = Pt(0)

    sub = doc.add_paragraph()
    sr = sub.add_run(f"{display_client(inc)} · {inc['report_date_long']}")
    sr.font.name = "Calibri"
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub.paragraph_format.space_after = Pt(2)

    sub2 = doc.add_paragraph()
    sr2 = sub2.add_run("La Vita Care, Community Services")
    sr2.font.name = "Calibri"
    sr2.font.size = Pt(10)
    sr2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    sub2.paragraph_format.space_after = Pt(6)

    # Horizontal rule via paragraph border
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '3D9B9C')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Status line
    status_p = doc.add_paragraph()
    sl = status_p.add_run(f"Status: {inc['status'].upper()}")
    sl.bold = True
    sl.font.size = Pt(11)
    sl.font.name = "Calibri"
    sl.font.color.rgb = RGBColor(0x3D, 0x9B, 0x9C)
    status_p.paragraph_format.space_after = Pt(14)

    def section_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(14)
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(0x3D, 0x9B, 0x9C)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(3)
        # Underline via border
        pPr2 = p._p.get_or_add_pPr()
        pBdr2 = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '7ACCC8')
        pBdr2.append(bot)
        pPr2.append(pBdr2)

    def add_section_table(rows_data):
        from docx.oxml.ns import qn as _qn
        t = doc.add_table(rows=0, cols=2)
        t.style = 'Table Grid'
        t.columns[0].width = Cm(5)
        t.columns[1].width = Cm(11.5)
        # Header row with brand teal
        hrow = t.add_row()
        hrow.cells[0].merge(hrow.cells[1])
        set_cell_bg(hrow.cells[0], "7ACCC8")
        hrow.cells[0].paragraphs[0].paragraph_format.space_before = Pt(1)
        hrow.cells[0].paragraphs[0].paragraph_format.space_after = Pt(1)
        for i, (label, value) in enumerate(rows_data):
            add_table_row(t, label, value, i)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    section_heading("Incident Overview")
    add_section_table([
        ("Report ID", inc["id"]),
        ("Date Submitted", inc["date_created"]),
        ("Event Date & Time", inc["event_datetime"]),
        ("Event Location", inc["event_location"]),
        ("Service", inc["service"]),
        ("Review Due Date", inc["review_due"]),
    ])

    section_heading("People Involved")
    add_section_table([
        ("Client", inc["client"]),
        ("Staff Member(s)", inc["staff"]),
        ("Other People Involved", inc["other_people"]),
        ("Incident Lead", inc["lead"]),
        ("Assigned To", inc["assignees"]),
    ])

    section_heading("Incident Details")
    add_section_table([
        ("(1.) Incident Details", inc["incident_details"]),
        ("(2.) Impact on Client", inc["impact"]),
        ("(3.) Immediate Actions Taken", inc["actions_taken"]),
        ("(4.) Required Follow-Up Actions", inc["follow_up"]),
    ])

    section_heading("Classification")
    add_section_table([
        ("Client Incident Type", inc["client_incident_type"]),
        ("WHS Incident Type", inc["whs_incident_type"]),
        ("Client Consulted", inc["client_consulted"]),
        ("Worker Consulted", inc["worker_consulted"]),
        ("Worker Injured", inc["worker_injured"]),
        ("NDIS Reportable", inc["ndis_reportable"]),
        ("WorkSafe Reportable", inc["worksafe_reportable"]),
    ])

    section_heading("Review & Compliance")
    add_section_table([
        ("Review Findings", inc["review_findings"]),
        ("Required Actions", inc["required_actions"]),
        ("Date LVC Review Done", inc["date_review_done"]),
        ("Date NDIS Notified", inc["date_ndis_notified"]),
        ("NDIS Reference No", inc["ndis_ref"]),
        ("Date Incident Closed", inc["date_closed"]),
    ])

    # Footer
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(16)
    pPr3 = footer_p._p.get_or_add_pPr()
    pBdr3 = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '3D9B9C')
    pBdr3.append(top)
    pPr3.append(pBdr3)
    fr = footer_p.add_run(f"Generated from ClickUp Incident Register · Task ID: {inc['id']} · La Vita Care")
    fr.font.size = Pt(8)
    fr.font.name = "Calibri"
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── UI ────────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="report-header">
  <h1>📋 Incident Report Generator</h1>
  <p>La Vita Care · Pull and export incident reports from ClickUp</p>
</div>
""", unsafe_allow_html=True)

api_key = get_api_key()
if not api_key:
    st.markdown('<div class="warning-box">⚠️ No API key configured in secrets. Enter your ClickUp API key below to continue.</div>', unsafe_allow_html=True)
    api_key = st.text_input("ClickUp API key:", type="password", placeholder="pk_XXXXXXXXXX_XXXXXXXXXXXXXXXXXXXX")

task_id = st.text_input("ClickUp Task ID", placeholder="e.g. 86d25kan1", help="Found in the URL when viewing a task in ClickUp")

if st.button("Generate Report", type="primary", disabled=not (task_id and api_key)):
    with st.spinner("Fetching incident from ClickUp..."):
        data, error = fetch_incident(task_id, api_key)

    if error:
        st.error(error)
    else:
        inc = parse_incident(data)
        st.success(f"✅ Incident loaded — {inc['client']} · {inc['event_datetime']}")

        # ── On-screen display ──
        st.markdown('<div class="section-title">Incident Overview</div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div class="field-row"><span class="field-label">Report ID</span><span class="field-value">{inc['id']}</span></div>
        <div class="field-row"><span class="field-label">Status</span><span class="field-value">{status_badge(inc['status'])}</span></div>
        <div class="field-row"><span class="field-label">Date Submitted</span><span class="field-value">{inc['date_created']}</span></div>
        <div class="field-row"><span class="field-label">Event Date & Time</span><span class="field-value">{inc['event_datetime']}</span></div>
        <div class="field-row"><span class="field-label">Event Location</span><span class="field-value">{inc['event_location']}</span></div>
        <div class="field-row"><span class="field-label">Service</span><span class="field-value">{inc['service']}</span></div>
        <div class="field-row"><span class="field-label">Review Due</span><span class="field-value">{inc['review_due']}</span></div>
        """, unsafe_allow_html=True)

        st.markdown('<div class="section-title">People Involved</div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div class="field-row"><span class="field-label">Client</span><span class="field-value">{inc['client']}</span></div>
        <div class="field-row"><span class="field-label">Staff Member(s)</span><span class="field-value">{inc['staff']}</span></div>
        <div class="field-row"><span class="field-label">Other People Involved</span><span class="field-value">{inc['other_people']}</span></div>
        <div class="field-row"><span class="field-label">Incident Lead</span><span class="field-value">{inc['lead']}</span></div>
        <div class="field-row"><span class="field-label">Assigned To</span><span class="field-value">{inc['assignees']}</span></div>
        """, unsafe_allow_html=True)

        st.markdown('<div class="section-title">Incident Details</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="narrative-label">(1.) Incident Details</div><div class="narrative-box">{inc["incident_details"]}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="narrative-label">(2.) Impact on Client</div><div class="narrative-box">{inc["impact"]}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="narrative-label">(3.) Immediate Actions Taken</div><div class="narrative-box">{inc["actions_taken"]}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="narrative-label">(4.) Required Follow-Up Actions</div><div class="narrative-box">{inc["follow_up"]}</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-title">Classification</div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div class="field-row"><span class="field-label">Client Incident Type</span><span class="field-value">{inc['client_incident_type']}</span></div>
        <div class="field-row"><span class="field-label">WHS Incident Type</span><span class="field-value">{inc['whs_incident_type']}</span></div>
        <div class="field-row"><span class="field-label">Client Consulted</span><span class="field-value">{inc['client_consulted']}</span></div>
        <div class="field-row"><span class="field-label">Worker Consulted</span><span class="field-value">{inc['worker_consulted']}</span></div>
        <div class="field-row"><span class="field-label">Worker Injured</span><span class="field-value">{inc['worker_injured']}</span></div>
        <div class="field-row"><span class="field-label">NDIS Reportable</span><span class="field-value">{inc['ndis_reportable']}</span></div>
        <div class="field-row"><span class="field-label">WorkSafe Reportable</span><span class="field-value">{inc['worksafe_reportable']}</span></div>
        """, unsafe_allow_html=True)

        st.markdown('<div class="section-title">Review & Compliance</div>', unsafe_allow_html=True)
        if inc["review_findings"] != "—":
            st.markdown(f'<div class="narrative-label">Review Findings</div><div class="narrative-box">{inc["review_findings"]}</div>', unsafe_allow_html=True)
        if inc["required_actions"] != "—":
            st.markdown(f'<div class="narrative-label">Required Actions</div><div class="narrative-box">{inc["required_actions"]}</div>', unsafe_allow_html=True)
        st.markdown(f"""
        <div class="field-row"><span class="field-label">Date LVC Review Done</span><span class="field-value">{inc['date_review_done']}</span></div>
        <div class="field-row"><span class="field-label">Date NDIS Notified</span><span class="field-value">{inc['date_ndis_notified']}</span></div>
        <div class="field-row"><span class="field-label">NDIS Reference No</span><span class="field-value">{inc['ndis_ref']}</span></div>
        <div class="field-row"><span class="field-label">Date Incident Closed</span><span class="field-value">{inc['date_closed']}</span></div>
        """, unsafe_allow_html=True)

        # ── Download ──
        st.divider()
        filename = build_filename(inc)
        word_buf = generate_word(inc)
        st.download_button(
            "⬇️ Download Word Report (.docx)",
            data=word_buf,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
